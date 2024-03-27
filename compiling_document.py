from docx import Document
from docx.shared import Inches
import os

def criar_documento(pasta, output):
    doc = Document()
    for arquivo_de_imagem in os.listdir('realmarcosdopampa'):
        if arquivo_de_imagem.lower().endswith('.jpg'):
            nome_da_imagem = os.path.splitext(arquivo_de_imagem)[0]
            doc.add_paragraph().add_run().add_picture(os.path.join(pasta, arquivo_de_imagem), width=Inches(4))
            text_file = os.path.join(pasta, nome_da_imagem + '.txt')
            if os.path.exists(text_file):
                with open(text_file, 'r', encoding='utf-8') as f:
                    legenda = f.read().strip()
                doc.add_paragraph(legenda)
            doc.add_paragraph().add_run('-' * 50).bold = True

    doc.save(output)

criar_documento('realmarcosdopampa', 'marcosdopampa.docx')